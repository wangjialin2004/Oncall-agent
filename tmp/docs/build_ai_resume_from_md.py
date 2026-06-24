import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path("output/doc/ai-resume.md")
OUT = Path("output/doc/王佳琳-AI应用-Java后端-简历.docx")
BLUE = "1F4D78"
LIGHT_BLUE = "B6C7D9"
INK = "111827"
MUTED = "4B5563"


def set_run_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size=10.5, bold=False, color=INK):
    font = style.font
    font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)


def add_border_bottom(paragraph, color="D1D5DB", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def para(doc, text="", before=0, after=4, line=1.08, align=None, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_inline_markdown(paragraph, text, base_size=10.0, color=INK):
    pos = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > pos:
            r = paragraph.add_run(text[pos : match.start()])
            set_run_font(r, size=base_size, color=color)
        r = paragraph.add_run(match.group(1))
        set_run_font(r, size=base_size, bold=True, color=color)
        pos = match.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, size=base_size, color=color)


def section_heading(doc, text):
    p = para(doc, before=7, after=3, line=1.0, keep_with_next=True)
    r = p.add_run(text)
    set_run_font(r, size=11.4, bold=True, color=BLUE)
    add_border_bottom(p, color=LIGHT_BLUE, size="6")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.48)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 1.04
    add_inline_markdown(p, text, base_size=9.35)


def clean_line(line):
    return line.strip().replace("  ", " ")


def build():
    raw_lines = SRC.read_text(encoding="utf-8").splitlines()
    lines = [line.rstrip() for line in raw_lines]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    set_style_font(doc.styles["Normal"], size=9.7)
    set_style_font(doc.styles["List Bullet"], size=9.35)

    idx = 0
    name = lines[idx].lstrip("# ").strip()
    p = para(doc, before=0, after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(name)
    set_run_font(r, size=20, bold=True, color=BLUE)
    idx += 1

    contact_parts = []
    while idx < len(lines):
        line = lines[idx].strip()
        idx += 1
        if line == "## 教育背景":
            idx -= 1
            break
        if not line:
            continue
        match = re.match(r"\*\*(.+?)：\*\*\s*(.+)", line)
        if match:
            label, value = match.groups()
            if label == "求职方向":
                p = para(doc, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
                r = p.add_run(value)
                set_run_font(r, size=10.2, bold=True, color=INK)
            elif label in {"电话", "邮箱", "学校", "专业"}:
                contact_parts.append(f"{label}: {value}")

    p = para(doc, after=7, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, part in enumerate(contact_parts):
        if i:
            sep = p.add_run("  |  ")
            set_run_font(sep, size=8.6, color="9CA3AF")
        r = p.add_run(part)
        set_run_font(r, size=8.6, color=MUTED)
    add_border_bottom(p, color=LIGHT_BLUE, size="8")

    current_section = None
    while idx < len(lines):
        line = lines[idx].strip()
        idx += 1
        if not line:
            continue

        if line.startswith("## "):
            current_section = line[3:].strip()
            section_heading(doc, current_section)
            continue

        if line.startswith("### "):
            p = para(doc, before=2, after=1, line=1.0, keep_with_next=True)
            r = p.add_run(line[4:].strip())
            set_run_font(r, size=10.2, bold=True, color=INK)
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue

        if line.startswith("**") and "：**" in line:
            p = para(doc, after=2, line=1.04)
            add_inline_markdown(p, line, base_size=9.35, color=MUTED)
            continue

        if line.startswith("**") and line.endswith("**"):
            p = para(doc, after=2, line=1.0, keep_with_next=True)
            add_inline_markdown(p, line, base_size=9.7)
            continue

        p = para(doc, after=3 if current_section != "教育背景" else 2, line=1.08)
        add_inline_markdown(p, clean_line(line), base_size=9.35)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
