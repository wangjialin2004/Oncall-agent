from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path("resume_template.docx")
OUT_DIR = Path("output/doc")
OUT = OUT_DIR / "\u738b\u4f73\u7433-\u540e\u7aef\u5f00\u53d1\u5de5\u7a0b\u5e08-\u6700\u7ec8\u7b80\u5386.docx"

DARK = "111827"


def set_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    rpr = run._element.get_or_add_rPr()
    run.font.name = name
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def fill_paragraph(
    paragraph,
    text,
    *,
    size=None,
    bold=None,
    color=DARK,
    left_cm=None,
    first_cm=None,
    before_pt=0,
    after_pt=0.8,
    line=1.0,
):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.line_spacing = line
    if left_cm is not None:
        paragraph.paragraph_format.left_indent = Cm(left_cm)
    if first_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_cm)


def header_line(paragraph, left, right, size=9.6):
    clear_paragraph(paragraph)
    for text in (left, "\t", right):
        run = paragraph.add_run(text)
        set_font(run, size=size, color=DARK)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0.6)
    paragraph.paragraph_format.line_spacing = 1.0


def role(paragraph, text, size=9.7):
    paragraph.style = doc.styles["Normal"]
    fill_paragraph(paragraph, text, size=size, bold=True, left_cm=0, first_cm=0, after_pt=0.3)


def body(paragraph, text, size=8.8):
    paragraph.style = doc.styles["Body Text"]
    fill_paragraph(paragraph, text, size=size, left_cm=0.05, first_cm=0, after_pt=0.7)


def bullet(paragraph, text, size=8.85):
    paragraph.style = doc.styles["List Paragraph"]
    fill_paragraph(
        paragraph,
        text,
        size=size,
        left_cm=0.64,
        first_cm=-0.26,
        after_pt=0.6,
        line=0.98,
    )


def blank(paragraph):
    fill_paragraph(paragraph, "", size=1, after_pt=0, line=1.0)


def replace_shape_texts(document):
    replacements = {
        "INTERNSHIP": "PROJECT",
        "EXPERIENCE": "EXPERIENCE",
        "\u5b9e\u4e60\u7ecf\u5386": "\u9879\u76ee\u7ecf\u5386",
        "CERTIFICATIONS": "TECH",
        "AND": "STACK",
        "AWARDS": "",
        "\u8bc1\u4e66\u5956\u52b1": "\u6280\u672f\u6808",
        "PROFESSIONAL": "PERSONAL",
        " SKILLS": " SUMMARY",
        "\u804c\u4e1a\u6280\u80fd": "\u4e2a\u4eba\u603b\u7ed3",
        "\u4e2a\u2f08\u7b80\u5386": "\u4e2a\u4eba\u7b80\u5386",
    }
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.text in replacements:
                run.text = replacements[run.text]
                set_font(run)


def patch_docx_text_literals(path):
    replacements = {
        "INTERNSHIP": "PROJECT",
        "\u5b9e\u4e60\u7ecf\u5386": "\u9879\u76ee\u7ecf\u5386",
        "CERTIFICATIONS": "TECH",
        "AND": "STACK",
        "AWARDS": "",
        "\u8bc1\u4e66\u5956\u52b1": "\u6280\u672f\u6808",
        "PROFESSIONAL": "PERSONAL",
        " SKILLS": " SUMMARY",
        "\u804c\u4e1a\u6280\u80fd": "\u4e2a\u4eba\u603b\u7ed3",
        "\u4e2a\u2f08\u7b80\u5386": "\u4e2a\u4eba\u7b80\u5386",
    }
    tmp = path.with_suffix(".patched.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                text = data.decode("utf-8")
                for old, new in replacements.items():
                    text = text.replace(f"<w:t>{old}</w:t>", f"<w:t>{new}</w:t>")
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


OUT_DIR.mkdir(parents=True, exist_ok=True)
doc = Document(SRC)

for section in doc.sections:
    section.start_type = WD_SECTION_START.CONTINUOUS
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0)

for style_name in ("Normal", "Body Text", "List Paragraph"):
    style = doc.styles[style_name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

replace_shape_texts(doc)
p = doc.paragraphs

fill_paragraph(
    p[2],
    "\u738b\u4f73\u7433\t/ \u5e94\u5c4a\u751f \u8f6f\u4ef6\u5de5\u7a0b",
    size=16.2,
    bold=True,
    after_pt=0.8,
)
header_line(
    p[3],
    "\u6c42\u804c\u610f\u5411\uff1a\u540e\u7aef\u5f00\u53d1\u5de5\u7a0b\u5e08 / AI \u5e94\u7528\u5f00\u53d1",
    "\u7535\u8bdd\uff1a18370092195",
    9.4,
)
header_line(
    p[4],
    "\u5b66\u6821\uff1a\u4e1c\u534e\u7406\u5de5\u5927\u5b66  \u4e13\u4e1a\uff1a\u8f6f\u4ef6\u5de5\u7a0b",
    "\u90ae\u7bb1\uff1awangde33@qq.com",
    9.4,
)
header_line(
    p[5],
    "\u9879\u76ee\u65b9\u5411\uff1aRAG / Agent / \u9ad8\u5e76\u53d1\u540e\u7aef",
    "GitHub\uff1agithub.com/wangjialin2004/Oncall-agent",
    9.0,
)

role(p[9], "\u4e1c\u534e\u7406\u5de5\u5927\u5b66 | \u8f6f\u4ef6\u5de5\u7a0b | \u672c\u79d1", 9.8)
role(
    p[10],
    "\u4e3b\u4fee\u8bfe\u7a0b\uff1a\u6570\u636e\u7ed3\u6784\u3001\u8ba1\u7b97\u673a\u7f51\u7edc\u3001\u8ba1\u7b97\u673a\u7ec4\u6210\u539f\u7406\u3001\u64cd\u4f5c\u7cfb\u7edf\u3001\u7ebf\u6027\u4ee3\u6570",
    8.8,
)
blank(p[11])
body(
    p[13],
    "\u5b66\u4e60\u91cd\u70b9\u8986\u76d6\u8ba1\u7b97\u673a\u57fa\u7840\u3001\u540e\u7aef\u5f00\u53d1\u4e0e\u6570\u636e\u7cfb\u7edf\uff1b\u8bfe\u5916\u9879\u76ee\u805a\u7126 Java \u540e\u7aef\u3001Python AI \u5e94\u7528\u4e0e\u5de5\u7a0b\u5316\u5b9e\u8df5\u3002",
    8.5,
)
for idx in (15, 16, 17, 19):
    blank(p[idx])

role(p[23], "\u667a\u80fd\u8fd0\u7ef4 OnCall Agent \u7cfb\u7edf | Python / FastAPI / Milvus / React / MCP", 9.6)
role(p[24], "\u4e2a\u4eba\u9879\u76ee / \u6838\u5fc3\u5f00\u53d1", 8.6)
body(
    p[25],
    "\u9762\u5411\u8fd0\u7ef4\u503c\u73ed\u544a\u8b66\u8bca\u65ad\uff0c\u63d0\u4f9b\u5bf9\u8bdd\u5165\u53e3\u3001SSE \u6d41\u5f0f\u8fc7\u7a0b\u5c55\u793a\u3001RAG \u77e5\u8bc6\u5e93\u68c0\u7d22\u3001\u7ecf\u9a8c\u8bb0\u5fc6\u4e0e\u5de5\u5177\u8c03\u7528\u95ed\u73af\u3002",
    8.45,
)
bullet(
    p[27],
    "\u8bbe\u8ba1 Agent \u8bca\u65ad\u4e3b\u6d41\u7a0b\uff0c\u5c06\u544a\u8b66\u5904\u7406\u62c6\u5206\u4e3a\u8ba1\u5212\u3001\u4e0a\u4e0b\u6587\u51c6\u5907\u3001\u4e13\u5bb6\u8def\u7531\u3001\u5de5\u5177\u6267\u884c\u3001\u8bc1\u636e\u6821\u9a8c\u548c\u7ed3\u679c\u6574\u7406\u3002",
)
bullet(
    p[28],
    "\u5b9e\u73b0\u77e5\u8bc6\u95ee\u7b54\u3001\u6307\u6807\u5206\u6790\u3001\u65e5\u5fd7\u5206\u6790\u3001\u53d8\u66f4\u6392\u67e5\u3001\u7efc\u5408\u8bca\u65ad 5 \u7c7b\u4e13\u5bb6\u6a21\u5757\uff0c\u7edf\u4e00\u8f93\u51fa\u8def\u7531\u3001\u5de5\u5177\u72b6\u6001\u548c\u9636\u6bb5\u6027\u7ed3\u8bba\u3002",
)
bullet(
    p[29],
    "\u642d\u5efa RAG \u77e5\u8bc6\u5e93\u94fe\u8def\uff0c\u652f\u6301\u6587\u6863\u4e0a\u4f20\u3001\u53d7\u4fe1\u76ee\u5f55\u7d22\u5f15\u3001\u6587\u672c\u5207\u5206\u3001Embedding\u3001Milvus \u5165\u5e93\u4e0e Top-K \u68c0\u7d22\u3002",
)
bullet(
    p[30],
    "\u5c06\u65e5\u5fd7\u67e5\u8be2\u548c Prometheus \u76d1\u63a7\u67e5\u8be2\u5c01\u88c5\u4e3a MCP \u5de5\u5177\uff0c\u5e76\u7528 pytest / Vitest \u8986\u76d6\u6838\u5fc3\u63a5\u53e3\u548c\u6d41\u5f0f\u4e8b\u4ef6\u3002",
)

role(p[33], "\u5546\u54c1\u79d2\u6740\u7cfb\u7edf | Java / Spring Boot / MySQL / Redis / Redisson / Lua", 9.6)
role(p[34], "\u8bfe\u7a0b / \u5b9e\u6218\u9879\u76ee", 8.6)
body(
    p[35],
    "\u4eff\u672c\u5730\u751f\u6d3b\u70b9\u8bc4\u4e1a\u52a1\u540e\u7aef\uff0c\u5305\u542b\u767b\u5f55\u3001\u5546\u94fa\u67e5\u8be2\u3001\u4f18\u60e0\u5238\u79d2\u6740\u3001\u535a\u5ba2\u4e92\u52a8\u3001\u5173\u6ce8\u3001\u7b7e\u5230\u4e0e\u9644\u8fd1\u5546\u94fa\u7b49\u6a21\u5757\u3002",
    8.45,
)
bullet(
    p[37],
    "\u57fa\u4e8e Redis \u5b9e\u73b0\u77ed\u4fe1\u9a8c\u8bc1\u7801\u767b\u5f55\u3001Token \u4f1a\u8bdd\u7ba1\u7406\u548c\u767b\u5f55\u6001\u5237\u65b0\uff0c\u964d\u4f4e\u9891\u7e41\u8bbf\u95ee\u6570\u636e\u5e93\u7684\u538b\u529b\u3002",
)
bullet(
    p[38],
    "\u8bbe\u8ba1\u5546\u94fa\u7f13\u5b58\u65b9\u6848\uff0c\u4f7f\u7528\u7f13\u5b58\u7a7a\u503c\u3001\u4e92\u65a5\u9501\u4e0e\u903b\u8f91\u8fc7\u671f\u5904\u7406\u7f13\u5b58\u7a7f\u900f\u548c\u70ed\u70b9\u7f13\u5b58\u51fb\u7a7f\u3002",
)
bullet(
    p[39],
    "\u7f16\u5199 Lua \u811a\u672c\u5b8c\u6210\u5e93\u5b58\u5224\u65ad\u3001\u91cd\u590d\u4e0b\u5355\u6821\u9a8c\u548c\u5e93\u5b58\u6263\u51cf\uff0c\u7ed3\u5408 Redis Stream \u4e0e Redisson \u5b9e\u73b0\u5f02\u6b65\u4e0b\u5355\u548c\u4e00\u4eba\u4e00\u5355\u3002",
)

body(p[41], "AI \u5e94\u7528\uff1aRAG\u3001Agent \u7f16\u6392\u3001\u5de5\u5177\u8c03\u7528\u3001Embedding\u3001Prompt \u8c03\u8bd5\u3001MCP\u3001SSE \u6d41\u5f0f\u8f93\u51fa\u3002", 8.8)
body(p[42], "\u540e\u7aef\u5f00\u53d1\uff1aPython\u3001FastAPI\u3001Pydantic\u3001Java\u3001Spring Boot\u3001MyBatis-Plus\u3001RESTful API\u3002", 8.8)
body(p[44], "\u6570\u636e\u4e0e\u4e2d\u95f4\u4ef6\uff1aMilvus\u3001MySQL\u3001SQLite\u3001Redis\u3001Redis Stream\u3001Redisson\u3001Lua\u3002", 8.8)
body(p[45], "\u524d\u7aef\u4e0e\u5de5\u5177\uff1aReact\u3001TypeScript\u3001Vite\u3001Git\u3001Docker\u3001Maven\u3001uv\u3001pytest\u3001Vitest\u3001Loguru\u3002", 8.8)

bullet(
    p[46],
    "\u5177\u5907 AI \u5e94\u7528\u548c\u540e\u7aef\u7cfb\u7edf\u5b8c\u6574\u9879\u76ee\u5b9e\u8df5\uff0c\u719f\u6089 RAG \u68c0\u7d22\u3001Agent \u5de5\u5177\u8c03\u7528\u4e0e\u5e38\u89c1\u540e\u7aef\u63a5\u53e3\u5f00\u53d1\u3002",
    8.75,
)
bullet(
    p[47],
    "\u540c\u65f6\u5177\u5907 Python AI \u5e94\u7528\u4e0e Java \u540e\u7aef\u5f00\u53d1\u7ecf\u9a8c\uff0c\u80fd\u5b8c\u6210\u63a5\u53e3\u8bbe\u8ba1\u3001\u8c03\u8bd5\u3001\u6d4b\u8bd5\u548c\u6587\u6863\u6574\u7406\u3002",
    8.75,
)

for idx in (0, 1, 7, 11, 15, 16, 17, 19, 21, 22, 31, 35, 40, 43):
    if idx < len(p) and not p[idx].text.strip():
        p[idx].paragraph_format.space_before = Pt(0)
        p[idx].paragraph_format.space_after = Pt(0)

doc.save(OUT)
patch_docx_text_literals(OUT)
print(OUT.resolve())
