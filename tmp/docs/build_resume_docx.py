from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "output/doc/AI_OnCall_Agent_Resume.docx"


def set_font(run, name="Arial", size=10.5, bold=False, color="111827"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", style=None, before=0, after=4, line=1.08, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run)
    return p


def add_bottom_border(paragraph, color="D1D5DB", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_style(style, font_name="Arial", east_asia="Microsoft YaHei", size=10.5, color="111827", bold=False):
    font = style.font
    font.name = font_name
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text):
    p = paragraph(doc, before=8, after=3, line=1.0)
    run = p.add_run(text)
    set_font(run, size=11.5, bold=True, color="0F766E")
    add_bottom_border(p, color="99F6E4", size="6")
    return p


def add_bullet(doc, text, after=2):
    p = paragraph(doc, after=after, line=1.08)
    p.style = doc.styles["List Bullet"]
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=9.8)
    return p


def add_role(doc, title, meta):
    p = paragraph(doc, before=4, after=1, line=1.0)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color="111827")
    if meta:
        p.add_run("  |  ")
        r2 = p.add_run(meta)
        set_font(r2, size=9.2, color="6B7280")
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.35)
section.bottom_margin = Cm(1.35)
section.left_margin = Cm(1.55)
section.right_margin = Cm(1.55)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

styles = doc.styles
set_style(styles["Normal"], size=10.0)
set_style(styles["List Bullet"], size=9.8)

title = paragraph(doc, after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
r = title.add_run("王佳林")
set_font(r, size=20, bold=True, color="111827")

target = paragraph(doc, after=3, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
r = target.add_run("AI 应用开发 / Python 后端开发 / AIOps Agent 方向")
set_font(r, size=10.5, color="0F766E", bold=True)

contact = paragraph(doc, after=6, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, part in enumerate(
    [
        "电话：[待补充]",
        "邮箱：[待补充]",
        "GitHub：github.com/wangjialin2004/Oncall-agent",
        "城市：[待补充]",
    ]
):
    if i:
        sep = contact.add_run("  |  ")
        set_font(sep, size=9.2, color="9CA3AF")
    rr = contact.add_run(part)
    set_font(rr, size=9.2, color="374151")
add_bottom_border(contact, color="E5E7EB", size="4")

add_heading(doc, "个人优势")
summary = (
    "具备 Python 后端、智能体编排、RAG 知识库与前端工程实践能力，独立完成面向 OnCall 运维场景的智能体平台。"
    "项目覆盖 FastAPI 服务、SSE 流式交互、多专家 Agent Harness、Milvus 向量检索、长期经验记忆、Prometheus/MCP 工具接入和 React 控制台，"
    "能够从业务问题拆解、系统设计、接口开发、测试验证到文档交付完整推进。"
)
p = paragraph(doc, after=4, line=1.15)
r = p.add_run(summary)
set_font(r, size=9.8)

add_heading(doc, "技术栈")
skills = [
    ("后端", "Python 3.11+, FastAPI, Pydantic, Uvicorn, SSE, Loguru, httpx/aiohttp"),
    ("AI / Agent", "OpenAI-compatible LLM, DashScope/Qwen, 多智能体 Harness, 工具调用, RAG, 提示词与评测"),
    ("数据与检索", "Milvus, SQLite, 向量索引, 文档切分, Embedding, 稠密/混合检索"),
    ("运维与集成", "Prometheus, MCP Server, Docker Compose, 环境变量配置, 服务健康检查"),
    ("前端与质量", "React 18, Vite, TypeScript, Vitest, Pytest, Ruff, Black, Pyright"),
]
for label, value in skills:
    p = paragraph(doc, after=2, line=1.08)
    r = p.add_run(f"{label}: ")
    set_font(r, size=9.8, bold=True, color="111827")
    r = p.add_run(value)
    set_font(r, size=9.8)

add_heading(doc, "项目经历")
add_role(doc, "智能 OnCall 运维平台 - Agent Gateway", "核心开发 | Python / FastAPI / React / Milvus / Prometheus | 2026")
project_intro = (
    "面向运维值班场景构建智能体平台，将告警输入自动转化为路由、规划、工具调用、证据自检和诊断结论，并通过 Web 控制台实时展示处理过程。"
)
p = paragraph(doc, after=3, line=1.12)
r = p.add_run(project_intro)
set_font(r, size=9.8)

bullets = [
    "设计统一助手入口 POST /api/assistant，使用 SSE 流式返回 route_selected、agent_event、tool_event、decision_event、content 等事件，提升诊断过程可观测性。",
    "实现多智能体 Harness 主循环，串联规划、上下文准备、专家路由、子任务执行、证据自检与降级收尾，支持知识库、指标、日志、变更和综合诊断专家协作。",
    "建设 RAG 知识库链路，支持文档上传、受信目录索引、文档切分、Embedding、Milvus 向量检索和检索评测脚本，沉淀可复用运维知识。",
    "实现长期记忆与服务基线能力，支持诊断经验、服务知识、指标基线和用户偏好的增删改查及向量索引重建。",
    "接入 Prometheus/Monitor 与 CLS MCP 服务，让 Agent 可调用监控指标、日志查询和服务知识工具形成证据闭环。",
    "搭建 React 18 + Vite 控制台，覆盖登录、会话历史、实时过程面板和服务基线管理，并使用 Vitest/Pytest 维护关键接口与前端交互测试。",
]
for b in bullets:
    add_bullet(doc, b)

add_role(doc, "本地 RAG 评测与调试流水线", "Python / Pytest / 向量检索")
for b in [
    "编写本地评测、RAG 样例生成、向量集合重建与调试脚本，便于定位检索召回、文档切分和工具调用链路中的问题。",
    "维护测试用例覆盖 Prometheus 集成、LLM 流式客户端、Harness 服务和本地评测流程，降低 Agent 行为改动带来的回归风险。",
]:
    add_bullet(doc, b)

add_heading(doc, "教育经历")
add_role(doc, "[学校名称待补充]", "[专业待补充] | [学历待补充] | [起止时间待补充]")
add_bullet(doc, "建议补充课程/竞赛/绩点/奖项中与后端开发、AI 应用、数据系统相关的内容。", after=4)

add_heading(doc, "实习 / 工作经历")
add_role(doc, "[公司或团队待补充]", "[岗位待补充] | [起止时间待补充]")
add_bullet(doc, "如暂无实习经历，可保留项目经历作为重点；如有经历，建议补充职责、技术栈、量化结果和业务影响。", after=4)

add_heading(doc, "补充信息")
for item in [
    "求职方向：AI 应用开发、Python 后端开发、AIOps/Agent 工程化方向。",
    "作品集：github.com/wangjialin2004/Oncall-agent；建议继续补充线上演示地址或项目截图链接。",
    "可根据目标岗位进一步调整关键词，例如 LLM Agent、RAG、FastAPI、Milvus、React、Prometheus、MCP。",
]:
    add_bullet(doc, item, after=2)

doc.save(OUT)
print(OUT)
